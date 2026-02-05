from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    document = Document()

    # Title
    title = document.add_heading('WOCC 2026 Call for Papers', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Navigation/Header info (simulated as text list for now as docx header is different)
    nav_para = document.add_paragraph()
    nav_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nav_items = [
        "Home", "About Us", "Organizing Committees", "Technical Program", "Speakers",
        "Paper Submission", "Call for Paper", "Call for Poster", "Registration",
        "Sponsorship", "Venue/Hotel", "IEEE Conference"
    ]
    nav_para.add_run(" | ".join(nav_items)).bold = True

    document.add_heading('Important Dates', level=1)
    dates_table = document.add_table(rows=1, cols=2)
    dates_table.style = 'Light Shading Accent 1'
    hdr_cells = dates_table.rows[0].cells
    hdr_cells[0].text = 'Event'
    hdr_cells[1].text = 'Date'
    
    dates_data = [
        ("Full Paper Submission deadline", "January 31, 2026 (extended)"),
        ("Notification of Acceptance", "March 14, 2026"),
        ("Camera-Ready", "April 4, 2026")
    ]
    
    for event, date in dates_data:
        row_cells = dates_table.add_row().cells
        row_cells[0].text = event
        row_cells[1].text = date

    document.add_heading('Submission Website and Guidelines', level=1)
    p = document.add_paragraph()
    p.add_run("All submission shall follow IEEE Templates, which are available at:\n")
    p.add_run("http://www.ieee.org/conferences_events/conferences/publishing/templates.html").font.color.rgb = RGBColor(0, 0, 255)
    
    p2 = document.add_paragraph()
    p2.add_run("Please submit your original papers at EDAS Website:\n")
    p2.add_run("https://edas.info/N34544").font.color.rgb = RGBColor(0, 0, 255)

    document.add_heading('Symposia (Tracks)', level=1)
    p_sym = document.add_paragraph("There are three symposias (tracks) of WOCC 2026:")
    tracks = [
        "Wireless Networks and Communications Symposium",
        "Optical Communications and Networks Symposium",
        "Artificial Intelligent and Big Data Analytics Symposium"
    ]
    for track in tracks:
        document.add_paragraph(track, style='List Bullet')
    
    document.add_paragraph("Please pick one of them to submit your manuscripts.")
    document.add_paragraph("4-5 page Full Papers are welcome. One or Two Best Papers will be awarded in each symposium.")

    document.add_heading('Important Reminder', level=1)
    reminder_text = (
        "At least one author of an accepted paper must register at the \"Regular\" or \"IEEE Members\" rates, "
        "even if the author is a student. A registration of \"Student\" category is not valid for covering the accepted paper. "
        "One \"Regular\" or \"IEEE Members\" registration can cover at most 2 accepted paper."
    )
    document.add_paragraph(reminder_text)

    document.add_heading('Submission Guidelines Detail', level=1)
    guidelines = (
        "All submissions should be written in English with a maximum paper length of five (5) printed pages (10-point font) "
        "including figures and tables (maximum 1 additional page with over length page charge of USD 150 if accepted). "
        "Papers exceeding 6 pages will not be accepted at EDAS. The submitted manuscripts should be in pdf format. "
        "All manuscripts are recommended to follow the IEEE templates."
    )
    document.add_paragraph(guidelines)
    
    document.add_paragraph("IEEE Templates are available at:")
    document.add_paragraph("http://www.ieee.org/conferences_events/conferences/publishing/templates.html")

    document.add_heading('Final Version Guidelines', level=1)
    document.add_paragraph("Camera-ready Submission Deadline: April 4th, 2026").bold = True
    
    final_para = document.add_paragraph(
        "To fit the IEEEXPLORE requirements your final manuscript needs to pass the verification of IEEE PDFeXpress Plus "
        "and IEEE Crosscheck Portal. You need to perform IEEE PDFeXpress Plus on your side. Please follow the instructions below! "
        "You will need to setup an account at IEEE PDFeXpress Plus and download the certificated pdf from your IEEE PDFeXpress Plus account. "
        "As for IEEE Crosscheck, you don’t need to worry about that we will perform the check for you afterwards."
    )

    document.add_heading('Before Creating a PDF', level=2)
    document.add_paragraph("1. Please remove subtitle such as (invited paper) for review purpose, as IEEEXPLORE does not capture subtitle.", style='List Number')
    document.add_paragraph("2. The IEEE e-Copyright Form must be completed on EDAS.", style='List Number')

    document.add_heading('Creating your PDF eXpress plus Account', level=2)
    document.add_paragraph("Instructions for generating camera-ready manuscript by using IEEE template and PDF eXpress:")
    document.add_paragraph("Access the IEEE PDF eXpress site: https://ieee-pdf-express.org/External/UsingIEEEPDFeXpress")
    
    document.add_heading('First-time users', level=3)
    document.add_paragraph("a. Click 'New Users - Click Here'.", style='List Bullet')
    document.add_paragraph("b. Enter '58016X' for the Conference ID, your email address, and choose a new password. Continue to enter information as prompted.", style='List Bullet')
    document.add_paragraph("c. You will receive online and email confirmation of successful account setup.", style='List Bullet')

    document.add_heading('Previous users (new conference)', level=3)
    document.add_paragraph("Enter 58016X for the Conference ID, your email address, and enter the password you used for your old account.")
    document.add_paragraph("a. When you click 'Login', you’ll receive an error saying you need to set up an account. Simply click 'Continue'. By entering your previously used email address and password combination, you will enable your old account for access to this new conference.", style='List Bullet')
    document.add_paragraph("b. Check that the contact information is still valid, and click 'Submit'.", style='List Bullet')
    document.add_paragraph("c. You will receive online and email confirmation of successful account setup.", style='List Bullet')

    document.add_heading('Returning users', level=3)
    document.add_paragraph("Enter 58016X for the Conference ID, email address and password.")
    document.add_paragraph("For each conference paper, click 'Create New Title'. Enter identifying text for the paper. Click 'Submit PDF for Checking' or 'Submit Source Files for Conversion'.")
    document.add_paragraph("Indicate platform, source file type (if applicable), click Browse and navigate to file, and click \"Upload File\". You will receive online and email confirmation of successful upload.")
    
    document.add_heading('Submit your paper to EDAS', level=2)
    document.add_paragraph("IEEE PDF eXpress converts the following file types to PDF: Rich Text Format, Freelance, (La)TeX, PageMaker, FrameMaker, QuarkXpress, Word Pro, Microsoft Word, WordPerfect.")

    filename = 'WOCC_2026_Call_For_Papers.docx'
    document.save(filename)
    print(f"Document saved as {filename}")

if __name__ == "__main__":
    create_document()
